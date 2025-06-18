import streamlit as st
from docx import Document
import re
from io import BytesIO
from collections import Counter

def extract_lines(doc):
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'^L(\d{1,3})(:)?\s*(.*)', text)
        if match:
            num = int(match.group(1))
            rest = match.group(3)
            lines.append((num, rest))
    return lines

def find_missing_and_duplicates(numbers):
    if not numbers:
        return [], []
    unique_numbers = sorted(set(numbers))
    full_range = list(range(unique_numbers[0], unique_numbers[-1] + 1))

    missing = [n for n in full_range if n not in numbers]
    duplicates = [f"L{n}" for n, c in Counter(numbers).items() if c > 1]

    return missing, duplicates

def fix_doc(lines, fix_missing=True, fix_duplicates=True):
    new_doc = Document()
    output_lines = []
    used = set()
    next_expected = lines[0][0]

    line_iter = iter(lines)

    while True:
        try:
            num, content = next(line_iter)

            # Insert missing lines
            while fix_missing and next_expected < num:
                output_lines.append((next_expected, "<< Missing Line >>"))
                next_expected += 1

            # Fix duplicates
            if fix_duplicates and num in used:
                # Assign the next available number
                while next_expected in used:
                    next_expected += 1
                output_lines.append((next_expected, content))
                used.add(next_expected)
                next_expected += 1
            else:
                output_lines.append((num, content))
                used.add(num)
                next_expected = num + 1

        except StopIteration:
            break

    # Write to docx
    for num, content in output_lines:
        new_doc.add_paragraph(f"L{num}: {content}")

    return new_doc

def generate_download_link(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("📄 Line Number Checker & Auto-Fix")
    st.write("Upload a `.docx` with `L{number}` lines. This app detects and optionally fixes skipped and duplicate numbers.")

    uploaded_file = st.file_uploader("Upload DOCX", type="docx")

    if uploaded_file:
        original_doc = Document(uploaded_file)
        lines = extract_lines(original_doc)
        line_numbers = [num for num, _ in lines]

        missing, duplicates = find_missing_and_duplicates(line_numbers)

        st.subheader("📋 Detected Line Numbers")
        st.write([f"L{n}" for n in line_numbers])

        if missing:
            st.subheader("❌ Missing Line Numbers")
            st.error(", ".join([f"L{m}" for m in missing]))
        else:
            st.success("✅ No missing line numbers!")

        if duplicates:
            st.subheader("⚠️ Duplicate Line Numbers")
            st.warning(", ".join(duplicates))
        else:
            st.success("✅ No duplicate line numbers!")

        fix_missing = st.checkbox("🛠️ Fix missing lines", value=True)
        fix_duplicates = st.checkbox("🔁 Fix duplicate lines", value=True)

        if fix_missing or fix_duplicates:
            if st.button("⚒️ Apply Fixes and Generate New File"):
                fixed_doc = fix_doc(lines, fix_missing=fix_missing, fix_duplicates=fix_duplicates)
                buffer = generate_download_link(fixed_doc)

                st.success("✅ Fixes applied.")
                st.download_button(
                    label="📥 Download Fixed DOCX",
                    data=buffer,
                    file_name="fixed_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
